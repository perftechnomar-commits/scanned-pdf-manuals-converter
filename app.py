import os
import tempfile

import pandas as pd
import streamlit as st
from docx import Document
from py_mistral_helper.MistralHelper import MistralHelper

st.set_page_config(page_title="OCR Verse", layout="centered")

st.title("📄 OCR Verse")

# Model selection
target_model = st.selectbox("Select OCR model", ["Mistral"], index=0)

# User input selection
option = st.radio("Choose input type", ["Document URL", "PDF", "Image URL", "Image"], horizontal=True)

# Input fields
document_url = None
document = None
image_url = None
image = None
extracted_text = st.session_state.get("extracted_text", "")
extracted_pages = st.session_state.get("extracted_pages", [])

if option == "Document URL":
    document_url = st.text_input("Enter document url")
elif option == "PDF":
    document = st.file_uploader("Upload a pdf", type=["pdf"])
elif option == "Image URL":
    image_url = st.text_input("Enter image url")
elif option == "Image":
    image = st.file_uploader("Upload an image", type=["png", "jpg", "jpeg"])

button_col1, space1, button_col2, space2, button_col3, space3, button_col4 = st.columns([1, 0.2, 1, 0.2, 1, 0.2, 1])

with button_col1:
    extract_btn = st.button("Extract text")

if extract_btn:
    mistral_helper = MistralHelper(api_key=st.secrets["MISTRAL_API_KEY"])
    extracted_pages = []

    if option == "Document URL" and document_url:
        _extracted_text = mistral_helper.extract_text_using_pdf_document_url(document_url)
        extracted_pages = [(i + 1, text.markdown) for i, text in enumerate(_extracted_text.pages)]

    elif document:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(document.read())
            temp_file_path = temp_file.name

        _extracted_text = mistral_helper.extract_text_using_pdf(temp_file_path)
        extracted_pages = [(i + 1, text.markdown) for i, text in enumerate(_extracted_text.pages)]

        os.remove(temp_file_path)

    elif image_url:
        _extracted_text = mistral_helper.extract_text_using_image_url(image_url)
        extracted_pages = [(i + 1, text.markdown) for i, text in enumerate(_extracted_text.pages)]

    elif image:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{image.type.split('/')[-1]}") as temp_file:
            temp_file.write(image.read())
            temp_file_path = temp_file.name

        _extracted_text = mistral_helper.extract_text_using_image_path(temp_file_path)
        extracted_pages = [(i + 1, text.markdown) for i, text in enumerate(_extracted_text.pages)]

        os.remove(temp_file_path)

    else:
        st.error("Please provide a valid input.")
        st.stop()

    if extracted_pages:
        extracted_text = "\n\n".join(f"**Page {page}**\n{text}" for page, text in extracted_pages)
        st.session_state["extracted_text"] = extracted_text
        st.session_state["extracted_pages"] = extracted_pages

if extracted_text:
    st.expander("Extracted Text", expanded=True).markdown(extracted_text)

    # Convert extracted text to DataFrame
    extracted_text_df = pd.DataFrame(extracted_pages, columns=["Page No", "Extracted Text"])


    def create_word_doc(filename):
        doc = Document()
        for page, content in extracted_pages:
            doc.add_heading(f"Page {page}", level=1)
            doc.add_paragraph(content)
        doc.save(filename)


    file_name = document.name.rsplit(".", 1)[0] if document else "output"

    with button_col2:
        csv_file = f"{file_name}.csv"
        st.download_button("📥 CSV", extracted_text_df.to_csv(index=False).encode("utf-8"), csv_file, "text/csv")
        if os.path.exists(csv_file):
            os.remove(csv_file)

    with button_col3:
        excel_file = f"{file_name}.xlsx"
        excel_buffer = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        extracted_text_df.to_excel(excel_buffer.name, index=False, engine="xlsxwriter")
        st.download_button("📥 Excel", excel_buffer.read(), excel_file,
                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        if os.path.exists(excel_file):
            os.remove(excel_file)

    with button_col4:
        word_file = f"{file_name}.docx"
        create_word_doc(word_file)
        with open(word_file, "rb") as f:
            st.download_button("📥 Word", f, word_file,
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        if os.path.exists(word_file):
            os.remove(word_file)
